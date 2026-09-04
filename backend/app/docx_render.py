"""将结构化 JSON 渲染为符合规范的 .docx。"""

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

CN_NUM = "一二三四五六七八九十"

TITLE_FONT = "黑体"
BODY_FONT = "宋体"


def _clean_xml_text(value: object) -> str:
    """移除 XML 1.0 不允许、因而无法写入 OOXML 的 Unicode 码点。"""
    text = str(value)
    return "".join(
        char
        for char in text
        if char in "\t\n\r"
        or "\x20" <= char <= "\ud7ff"
        or "\ue000" <= char <= "\ufffd"
        or "\U00010000" <= char <= "\U0010ffff"
    )


def _new_doc():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    return doc


def _set_font(run, name, size, bold=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)


def _add_paragraph(doc, text, size=12, bold=False, center=False, font=BODY_FONT, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if indent:
        p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(_clean_xml_text(text))
    _set_font(run, font, size, bold)
    return p


def _cn(i):
    return CN_NUM[i - 1] if 1 <= i <= len(CN_NUM) else str(i)


def render_report(data: dict) -> Document:
    report = data.get("report") or {}
    doc = _new_doc()
    _add_paragraph(doc, report.get("title") or "工作汇报", size=16, bold=True, center=True, font=TITLE_FONT)
    for i, section in enumerate(report.get("sections") or [], start=1):
        category = section.get("category") or "其他"
        _add_paragraph(doc, f"{_cn(i)}、{category}", size=14, bold=True, font=TITLE_FONT)
        for item in section.get("items") or []:
            summary = item.get("summary") or ""
            if summary:
                _add_paragraph(doc, summary, size=12, bold=True, font=TITLE_FONT)
            for label, key in (("日期", "date"), ("进展", "detail"), ("结果", "result"), ("下周计划", "next_step")):
                value = (item.get(key) or "").strip()
                if value:
                    _add_paragraph(doc, f"{label}：{value}", size=12, indent=True)
    return doc


def render_tech(data: dict) -> Document:
    tech = data.get("tech_summary") or {}
    doc = _new_doc()
    _add_paragraph(doc, tech.get("title") or "技术总结", size=16, bold=True, center=True, font=TITLE_FONT)
    topics = tech.get("topics") or []
    if not topics:
        _add_paragraph(doc, "本周无明确技术内容。", size=12, indent=True)
    for i, topic in enumerate(topics, start=1):
        _add_paragraph(doc, f"{_cn(i)}、{topic.get('topic') or '未命名主题'}", size=14, bold=True, font=TITLE_FONT)
        explanation = (topic.get("explanation") or "").strip()
        if explanation:
            _add_paragraph(doc, f"简介：{explanation}", size=12, indent=True)
        key_points = topic.get("key_points") or []
        if key_points:
            _add_paragraph(doc, "关键技术点：", size=12, bold=True)
            for kp in key_points:
                _add_paragraph(doc, f"• {kp}", size=12)
        related = topic.get("related_items") or []
        if related:
            _add_paragraph(doc, f"与本周工作的关联：{'；'.join(related)}", size=12, indent=True)
        references = topic.get("references") or []
        if references:
            _add_paragraph(doc, "参考资料：", size=12, bold=True)
            for ref in references:
                _add_paragraph(doc, f"• {ref}", size=12)
    return doc


def render_custom(definition: dict, data: dict) -> Document:
    """按受控模板定义渲染一份自定义 Word 文档。"""
    doc = _new_doc()
    _add_paragraph(
        doc,
        data.get("title") or "自定义汇报",
        size=16,
        bold=True,
        center=True,
        font=TITLE_FONT,
    )
    section_values = {section.get("id"): section for section in data.get("sections") or []}
    for index, section in enumerate(definition.get("sections") or [], start=1):
        _add_paragraph(
            doc,
            f"{_cn(index)}、{section.get('title') or '未命名章节'}",
            size=14,
            bold=True,
            font=TITLE_FONT,
        )
        supplied = section_values.get(section.get("id")) or {}
        block_values = {block.get("id"): block for block in supplied.get("blocks") or []}
        for block in section.get("blocks") or []:
            value = block_values.get(block.get("id")) or {}
            block_type = block.get("type")
            label = block.get("label") or "内容"
            if block_type == "paragraph":
                text = (value.get("text") or "").strip()
                if text:
                    _add_paragraph(doc, label, size=12, bold=True, font=TITLE_FONT)
                    _add_paragraph(doc, text, size=12, indent=True)
            elif block_type == "field":
                text = (value.get("text") or "").strip()
                if text:
                    _add_paragraph(doc, f"{label}：{text}", size=12, indent=True)
            elif block_type in {"bullet_list", "numbered_list"}:
                items = value.get("items") or []
                if items:
                    _add_paragraph(doc, label, size=12, bold=True, font=TITLE_FONT)
                    for item_index, item in enumerate(items, start=1):
                        prefix = "•" if block_type == "bullet_list" else f"{item_index}."
                        _add_paragraph(doc, f"{prefix} {item}", size=12)
            elif block_type == "table":
                rows = value.get("rows") or []
                if not rows:
                    continue
                columns = block.get("columns") or []
                _add_paragraph(doc, label, size=12, bold=True, font=TITLE_FONT)
                table = doc.add_table(rows=1, cols=len(columns))
                table.style = "Table Grid"
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                for column_index, column in enumerate(columns):
                    cell = table.rows[0].cells[column_index]
                    cell.text = ""
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                    run = cell.paragraphs[0].add_run(_clean_xml_text(column.get("label") or "列"))
                    _set_font(run, TITLE_FONT, 11, True)
                for row in rows:
                    cells = table.add_row().cells
                    for column_index, column in enumerate(columns):
                        cells[column_index].text = ""
                        cells[column_index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        run = cells[column_index].paragraphs[0].add_run(_clean_xml_text(row.get(column.get("id"), "")))
                        _set_font(run, BODY_FONT, 11)
                doc.add_paragraph()
    return doc
