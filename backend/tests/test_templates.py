import io
import json
import zipfile
from dataclasses import replace

import pytest
from docx import Document
from fastapi.testclient import TestClient

from app import attachments, storage, template_system, user_settings, visitor


@pytest.fixture(autouse=True)
def isolated_runtime_db(tmp_path, monkeypatch):
    monkeypatch.setattr(storage, "DB_PATH", tmp_path / "template-runtime.db")
    storage.init_db()
    yield


@pytest.fixture()
def client(tmp_path, monkeypatch):
    monkeypatch.setattr(storage, "DB_PATH", tmp_path / "templates.db")
    storage.init_db()
    template_system.clear_custom_conversations()
    from app.main import app

    return TestClient(app)


def _save_settings(client, **overrides):
    payload = {
        "week_one_start": user_settings.current_monday().isoformat(),
        "purpose_mode": overrides.get("purpose_mode", "default"),
        "custom_purpose_name": overrides.get("custom_purpose_name", ""),
        "custom_purpose_description": overrides.get("custom_purpose_description", ""),
        "detail_level": "standard",
        "tone": "natural",
    }
    return client.put("/api/settings", json=payload)


def _owner_id(client):
    client.get("/")
    token = visitor._validate_cookie(client.cookies.get(visitor.COOKIE_NAME))
    assert token
    return visitor.visitor_id(token)


def _create_template(client, name="科研模板"):
    draft = client.post("/api/template-drafts", json={"source_type": "manual"}).json()
    saved = client.post(
        "/api/templates",
        json={"draft_id": draft["id"], "draft_revision": draft["revision"], "name": name},
    )
    assert saved.status_code == 200
    return saved.json()


def test_template_definition_and_custom_document_validation():
    definition = template_system.default_definition()
    title = template_system.render_title(
        definition,
        user_settings.current_monday(),
        user_settings.current_monday(),
    )
    preview = template_system.placeholder_document(definition, title)
    validated = template_system.validate_custom_document(definition, preview, title)
    assert validated["title"] == title
    assert validated["sections"][0]["blocks"][0]["id"] == "main_content"

    preview["sections"][0]["blocks"][0]["text"] = ""
    with pytest.raises(ValueError, match="必填内容"):
        template_system.validate_custom_document(definition, preview, title)


def test_legacy_custom_setting_migrates_once_to_draft(tmp_path, monkeypatch):
    monkeypatch.setattr(storage, "DB_PATH", tmp_path / "legacy-template.db")
    storage.init_db()
    current = user_settings.current_monday().isoformat()
    storage.save_settings(
        "owner",
        {
            "week_one_start": current,
            "purpose_mode": "custom",
            "custom_purpose_name": "旧科研周报",
            "custom_purpose_description": "重点记录实验设计、结果和下周计划。",
            "detail_level": "standard",
            "tone": "natural",
            "onboarding_completed": True,
        },
    )

    storage.init_db()
    storage.init_db()

    rows = storage.list_templates("owner")
    assert len(rows) == 1
    assert rows[0]["status"] == "draft"
    assert rows[0]["source_type"] == "legacy"
    assert storage.get_settings("owner")["purpose_mode"] == "default"
    definition = json.loads(rows[0]["definition_json"])
    assert "实验设计" in definition["sections"][0]["blocks"][0]["instruction"]


def test_template_crud_selection_and_reserved_names(client):
    assert _save_settings(client).status_code == 200
    saved = _create_template(client)
    listed = client.get("/api/templates").json()
    assert listed["selected_template_id"] == saved["id"]
    assert [item["name"] for item in listed["templates"]] == ["科研模板"]

    renamed = client.patch(f"/api/templates/{saved['id']}", json={"name": "项目模板"})
    assert renamed.status_code == 200
    assert renamed.json()["name"] == "项目模板"

    duplicate_draft = client.post("/api/template-drafts", json={"source_type": "manual"}).json()
    duplicate = client.post(
        "/api/templates",
        json={
            "draft_id": duplicate_draft["id"],
            "draft_revision": duplicate_draft["revision"],
            "name": "项目模板",
        },
    )
    assert duplicate.status_code == 400
    reserved = client.patch(f"/api/templates/{saved['id']}", json={"name": "自定义"})
    assert reserved.status_code == 400

    deleted = client.delete(f"/api/templates/{saved['id']}")
    assert deleted.status_code == 200
    assert storage.get_settings(_owner_id(client))["selected_template_id"] is None


def test_template_draft_revision_rejects_stale_update(client):
    assert _save_settings(client).status_code == 200
    draft = client.post("/api/template-drafts", json={"source_type": "manual"}).json()
    assert draft["revision"] == 0

    definition = draft["definition"]
    definition["sections"][0]["title"] = "第一次修改"
    first = client.put(
        f"/api/template-drafts/{draft['id']}",
        json={"definition": definition, "base_revision": 0},
    )
    assert first.status_code == 200
    assert first.json()["revision"] == 1

    definition["sections"][0]["title"] = "过期请求"
    stale = client.put(
        f"/api/template-drafts/{draft['id']}",
        json={"definition": definition, "base_revision": 0},
    )
    assert stale.status_code == 409
    assert "已在其他页面或请求中更新" in stale.json()["detail"]


def test_template_save_rejects_stale_draft_revision_without_discarding_latest(client):
    assert _save_settings(client).status_code == 200
    draft = client.post("/api/template-drafts", json={"source_type": "manual"}).json()
    definition = draft["definition"]
    definition["sections"][0]["title"] = "最新草稿"
    updated = client.put(
        f"/api/template-drafts/{draft['id']}",
        json={"definition": definition, "base_revision": draft["revision"]},
    ).json()

    stale = client.post(
        "/api/templates",
        json={
            "draft_id": draft["id"],
            "draft_revision": draft["revision"],
            "name": "不会保存",
        },
    )

    assert stale.status_code == 409
    current = template_system.get_draft(_owner_id(client), draft["id"])
    assert current and current.revision == updated["revision"]
    assert storage.list_templates(_owner_id(client)) == []


def test_two_edit_drafts_cannot_silently_overwrite_formal_template(client):
    assert _save_settings(client).status_code == 200
    saved = _create_template(client, "并发模板")
    first = client.post(f"/api/templates/{saved['id']}/edit-draft").json()
    second = client.post(f"/api/templates/{saved['id']}/edit-draft").json()

    first_definition = first["definition"]
    first_definition["sections"][0]["title"] = "先保存的版本"
    first = client.put(
        f"/api/template-drafts/{first['id']}",
        json={"definition": first_definition, "base_revision": first["revision"]},
    ).json()
    second_definition = second["definition"]
    second_definition["sections"][0]["title"] = "不应覆盖的版本"
    second = client.put(
        f"/api/template-drafts/{second['id']}",
        json={"definition": second_definition, "base_revision": second["revision"]},
    ).json()

    first_save = client.put(
        f"/api/templates/{saved['id']}",
        json={
            "draft_id": first["id"],
            "draft_revision": first["revision"],
            "name": saved["name"],
        },
    )
    stale_save = client.put(
        f"/api/templates/{saved['id']}",
        json={
            "draft_id": second["id"],
            "draft_revision": second["revision"],
            "name": saved["name"],
        },
    )

    assert first_save.status_code == 200
    assert stale_save.status_code == 409
    current = client.get(f"/api/templates/{saved['id']}").json()
    assert current["definition"]["sections"][0]["title"] == "先保存的版本"
    assert template_system.get_draft(_owner_id(client), second["id"]) is not None


def test_templates_are_isolated_by_visitor(client):
    _save_settings(client)
    saved = _create_template(client)

    stranger = TestClient(client.app)
    _save_settings(stranger)
    assert stranger.get(f"/api/templates/{saved['id']}").status_code == 404
    assert stranger.delete(f"/api/templates/{saved['id']}").status_code == 404


def test_chat_session_cannot_switch_between_weekly_and_custom(client, monkeypatch):
    from app import main as main_module

    _save_settings(client)
    saved = _create_template(client)
    monkeypatch.setattr(
        main_module.app.state,
        "settings",
        replace(main_module.app.state.settings, deepseek_api_key="fake-key"),
    )
    template_system.clear_custom_conversations(_owner_id(client))

    class EmptyCompletions:
        def create(self, **kwargs):
            return []

    class FakeOpenAI:
        def __init__(self, **kwargs):
            self.chat = type("Chat", (), {"completions": EmptyCompletions()})()

    from app import agent

    monkeypatch.setattr(agent, "OpenAI", FakeOpenAI)
    first = client.post(
        "/api/chat",
        json={"session_id": "fixed-session", "message": "先使用默认周报", "template_id": None},
    )
    assert first.status_code == 200
    switched = client.post(
        "/api/chat",
        json={
            "session_id": "fixed-session",
            "message": "再换模板",
            "template_id": saved["id"],
        },
    )
    assert switched.status_code == 409


def test_template_learning_attachment_context_preserves_docx_structure():
    document = Document()
    document.add_heading("项目进展", level=1)
    document.add_paragraph("完成模板系统")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "任务"
    table.cell(0, 1).text = "状态"
    table.cell(1, 0).text = "接口"
    table.cell(1, 1).text = "完成"
    buffer = io.BytesIO()
    document.save(buffer)
    record = attachments.add("owner", "template-session", "sample.docx", "application/docx", buffer.getvalue())

    context = attachments.template_context_for("owner", "template-session", [record["id"]])
    assert "项目进展" in context
    assert "[表格 1 开始]" in context
    assert "任务 | 状态" in context

    with pytest.raises(attachments.AttachmentError, match="1–5"):
        attachments.template_context_for("owner", "template-session", [])
    with pytest.raises(attachments.AttachmentError, match="1–5"):
        attachments.template_context_for("owner", "template-session", [record["id"]] * 6)


def test_template_learning_rejects_unsupported_attachment():
    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w") as archive:
        archive.writestr("note.txt", "content")
    record = attachments.add("owner", "image-session", "sample.zip", "application/zip", buffer.getvalue())
    with pytest.raises(attachments.AttachmentError, match="仅支持"):
        attachments.template_context_for("owner", "image-session", [record["id"]])


def test_analyze_samples_returns_incompatible_without_creating_draft(monkeypatch):
    record = attachments.add("owner", "analysis", "one.txt", "text/plain", b"heading: value")

    class Completions:
        def create(self, **kwargs):
            return type(
                "Response",
                (),
                {
                    "choices": [
                        type(
                            "Choice",
                            (),
                            {
                                "message": type(
                                    "Message",
                                    (),
                                    {
                                        "content": json.dumps(
                                            {
                                                "status": "incompatible",
                                                "reason": "结构没有共同点",
                                                "warnings": [],
                                                "definition": None,
                                            },
                                            ensure_ascii=False,
                                        )
                                    },
                                )()
                            },
                        )()
                    ]
                },
            )()

    class FakeOpenAI:
        def __init__(self, **kwargs):
            self.chat = type("Chat", (), {"completions": Completions()})()

    monkeypatch.setattr(template_system, "OpenAI", FakeOpenAI)
    result = template_system.analyze_samples("owner", "analysis", [record["id"]], "key", "https://invalid", "model")
    assert result == {"status": "incompatible", "reason": "结构没有共同点", "warnings": []}


def test_custom_chat_archives_snapshot_and_exports_after_template_delete(client, monkeypatch):
    from app import main as main_module

    _save_settings(client)
    saved = _create_template(client, "实验周报")
    definition = saved["definition"]
    output = {
        "title": "由服务端覆盖",
        "sections": [
            {
                "id": definition["sections"][0]["id"],
                "title": definition["sections"][0]["title"],
                "blocks": [
                    {
                        "id": definition["sections"][0]["blocks"][0]["id"],
                        "type": "paragraph",
                        "text": "完成自定义模板联调",
                        "items": [],
                        "rows": [],
                    }
                ],
            }
        ],
    }
    content = "已整理完成。\n" + template_system.CUSTOM_FINAL_MARKER + "\n" + json.dumps(output, ensure_ascii=False)

    class Chunk:
        choices = [type("Choice", (), {"delta": type("Delta", (), {"content": content})()})()]

    class Completions:
        def create(self, **kwargs):
            return [Chunk()]

    class FakeOpenAI:
        def __init__(self, **kwargs):
            self.chat = type("Chat", (), {"completions": Completions()})()

    monkeypatch.setattr(template_system, "OpenAI", FakeOpenAI)
    monkeypatch.setattr(
        main_module.app.state,
        "settings",
        replace(main_module.app.state.settings, deepseek_api_key="fake-key"),
    )
    response = client.post(
        "/api/chat",
        json={
            "session_id": "custom-chat",
            "message": "本周完成自定义模板联调，生成。",
            "attachment_ids": [],
            "mode": "advanced",
            "template_id": saved["id"],
        },
    )
    assert response.status_code == 200
    events = [json.loads(line[6:]) for line in response.text.splitlines() if line.startswith("data: ")]
    final = next(event for event in events if event["type"] == "final")
    assert final["output_kind"] == "custom"
    week_id = final["week_id"]

    assert client.delete(f"/api/templates/{saved['id']}").status_code == 200
    history = client.get(f"/api/weeks/{week_id}").json()
    assert history["output_kind"] == "custom"
    assert history["template_name"] == "实验周报"
    assert history["document"]["sections"][0]["blocks"][0]["text"] == "完成自定义模板联调"
    exported = client.get(f"/api/weeks/{week_id}/export?doc=custom")
    assert exported.status_code == 200
    assert exported.content[:2] == b"PK"


def test_active_custom_chat_keeps_original_template_snapshot(client, monkeypatch):
    from app import main as main_module

    _save_settings(client)
    saved = _create_template(client, "会话快照模板")
    original = saved["definition"]
    final_document = template_system.placeholder_document(original)
    final_document["sections"][0]["blocks"][0]["text"] = "保留旧模板结构"
    responses = [
        "请再补充一个结果。",
        "已整理。" + template_system.CUSTOM_FINAL_MARKER + json.dumps(final_document, ensure_ascii=False),
    ]

    class Chunk:
        def __init__(self, content):
            self.choices = [type("Choice", (), {"delta": type("Delta", (), {"content": content})()})()]

    class Completions:
        def create(self, **kwargs):
            return [Chunk(responses.pop(0))]

    class FakeOpenAI:
        def __init__(self, **kwargs):
            self.chat = type("Chat", (), {"completions": Completions()})()

    monkeypatch.setattr(template_system, "OpenAI", FakeOpenAI)
    monkeypatch.setattr(
        main_module.app.state,
        "settings",
        replace(main_module.app.state.settings, deepseek_api_key="fake-key"),
    )
    first = client.post(
        "/api/chat",
        json={
            "session_id": "snapshot-chat",
            "message": "本周完成了模板功能。",
            "template_id": saved["id"],
        },
    )
    assert first.status_code == 200
    assert '"type": "final"' not in first.text

    edit = client.post(f"/api/templates/{saved['id']}/edit-draft").json()
    changed = edit["definition"]
    changed["sections"][0]["id"] = "new_section"
    changed["sections"][0]["blocks"][0]["id"] = "new_block"
    updated = client.put(
        f"/api/template-drafts/{edit['id']}",
        json={"definition": changed, "base_revision": edit["revision"]},
    )
    assert updated.status_code == 200
    assert (
        client.put(
            f"/api/templates/{saved['id']}",
            json={
                "draft_id": edit["id"],
                "draft_revision": updated.json()["revision"],
                "name": saved["name"],
            },
        ).status_code
        == 200
    )

    second = client.post(
        "/api/chat",
        json={
            "session_id": "snapshot-chat",
            "message": "结果正常，生成。",
            "template_id": saved["id"],
        },
    )
    events = [json.loads(line[6:]) for line in second.text.splitlines() if line.startswith("data: ")]
    final = next(event for event in events if event["type"] == "final")
    archived = client.get(f"/api/weeks/{final['week_id']}").json()
    assert archived["definition"]["sections"][0]["id"] == original["sections"][0]["id"]
    assert archived["document"]["sections"][0]["blocks"][0]["text"] == "保留旧模板结构"
