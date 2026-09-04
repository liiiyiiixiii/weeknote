import io

from docx import Document
from docx.oxml.ns import qn

from app import docx_render

SAMPLE = {
    "report": {
        "title": "第 33 周工作汇报（2025.08.11–2025.08.17）",
        "sections": [
            {
                "category": "工作",
                "items": [
                    {
                        "date": "2025-08-12",
                        "summary": "完成了XX模块",
                        "detail": "实现XX功能",
                        "result": "已上线",
                        "next_step": "下周联调",
                    }
                ],
            }
        ],
    },
    "tech_summary": {
        "title": "第 33 周技术总结（2025.08.11–2025.08.17）",
        "topics": [
            {
                "topic": "FastAPI",
                "related_items": ["完成了XX模块"],
                "explanation": "后端框架",
                "key_points": ["路由", "依赖注入"],
                "references": ["https://fastapi.tiangolo.com"],
            }
        ],
    },
}


def _save(doc):
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Document(buf)


def test_report_render_contains_title_and_sections():
    doc = _save(docx_render.render_report(SAMPLE))
    texts = [p.text for p in doc.paragraphs]
    assert texts[0] == "第 33 周工作汇报（2025.08.11–2025.08.17）"
    assert any("一、工作" in t for t in texts)
    assert any("完成了XX模块" in t for t in texts)
    assert any(t.startswith("结果：已上线") for t in texts)


def test_tech_render_contains_topics():
    doc = _save(docx_render.render_tech(SAMPLE))
    texts = [p.text for p in doc.paragraphs]
    assert any("一、FastAPI" in t for t in texts)
    assert any("关键技术点" in t for t in texts)
    assert any("与本周工作的关联" in t for t in texts)


def test_title_font_is_heiti():
    doc = _save(docx_render.render_report(SAMPLE))
    run = doc.paragraphs[0].runs[0]
    assert run.font.size.pt == 16
    assert run.font.bold is True
    assert run._element.rPr.rFonts.get(qn("w:eastAsia")) == "黑体"


def test_report_and_tech_strip_xml_10_invalid_characters():
    data = {
        "report": {
            "title": "周\x00报\x0b标题",
            "sections": [{"category": "工\ud800作", "items": [{"summary": "完成\x01事项"}]}],
        },
        "tech_summary": {
            "title": "技术\x0c总结",
            "topics": [{"topic": "解析\ufffe", "key_points": ["安全\x02写入"]}],
        },
    }

    report = _save(docx_render.render_report(data))
    tech = _save(docx_render.render_tech(data))

    assert "".join(p.text for p in report.paragraphs) == "周报标题一、工作完成事项"
    assert "技术总结" in "".join(p.text for p in tech.paragraphs)
    assert "安全写入" in "".join(p.text for p in tech.paragraphs)


def test_custom_document_strips_invalid_characters_from_all_block_types():
    definition = {
        "sections": [
            {
                "id": "s",
                "title": "章\x0b节",
                "blocks": [
                    {"id": "p", "type": "paragraph", "label": "段\x01落"},
                    {"id": "f", "type": "field", "label": "字\x02段"},
                    {"id": "l", "type": "bullet_list", "label": "列\x03表"},
                    {
                        "id": "t",
                        "type": "table",
                        "label": "表\x04格",
                        "columns": [{"id": "c", "label": "列\x05名"}],
                    },
                ],
            }
        ]
    }
    data = {
        "title": "自\x00定义",
        "sections": [
            {
                "id": "s",
                "blocks": [
                    {"id": "p", "text": "正\x06文"},
                    {"id": "f", "text": "值\x07"},
                    {"id": "l", "items": ["项\x08目"]},
                    {"id": "t", "rows": [{"c": "单\x0b元格"}]},
                ],
            }
        ],
    }

    doc = _save(docx_render.render_custom(definition, data))
    paragraph_text = "".join(p.text for p in doc.paragraphs)
    table_text = "".join(cell.text for table in doc.tables for row in table.rows for cell in row.cells)

    assert "自定义" in paragraph_text
    assert "章节" in paragraph_text
    assert "段落" in paragraph_text and "正文" in paragraph_text
    assert "字段：值" in paragraph_text
    assert "列表" in paragraph_text and "项目" in paragraph_text
    assert table_text == "列名单元格"
